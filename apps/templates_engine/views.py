from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.exceptions import ValidationError
from django.db.models import F
from rest_framework import permissions
from .models import DocumentTemplate, TemplateUsage
from .serializers import DocumentTemplateSerializer
from .tasks import generate_document_from_template_sync


def extract_placeholders(file):
    """
    Scan a DOCX/XLSX for {{key}} patterns and return the unique keys.
    Also handles placeholders that may be split across multiple runs
    by joining paragraph text first.
    """
    import re
    pattern = re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}")
    found = set()

    if file.name.endswith((".docx", ".doc")):
        from docx import Document
        doc = Document(file)
        # Use full paragraph text (handles split runs)
        for para in doc.paragraphs:
            found.update(pattern.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found.update(pattern.findall(para.text))
        # Also scan text boxes and headers/footers if present
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf:
                    for para in hf.paragraphs:
                        found.update(pattern.findall(para.text))

    elif file.name.endswith((".xlsx", ".xls")):
        import openpyxl
        wb = openpyxl.load_workbook(file)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if isinstance(cell, str):
                        found.update(pattern.findall(cell))

    return sorted(found)


class DocumentTemplateViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentTemplateSerializer
    queryset = DocumentTemplate.objects.all()

    def get_permissions(self):
        if self.action in ("create", "update", "partial_update", "destroy", "duplicate"):
            return [permissions.IsAuthenticated(), permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    def get_queryset(self):
        queryset = super().get_queryset()
        if self.action == "list":
            queryset = queryset.filter(is_active=True)
            # Optional filters
            category = self.request.query_params.get("category")
            document_type_id = (
                self.request.query_params.get("document_type")
                or self.request.query_params.get("document_type_id")
            )
            type_ = self.request.query_params.get("type")
            search = self.request.query_params.get("search")
            if category:
                queryset = queryset.filter(category=category)
            if document_type_id:
                queryset = queryset.filter(document_type_id=document_type_id)
            if type_:
                queryset = queryset.filter(type=type_)
            if search:
                queryset = queryset.filter(name__icontains=search)
        return queryset.select_related("created_by", "document_type")

    def perform_create(self, serializer):
        file = self.request.FILES.get("file")
        placeholders = []

        if file:
            try:
                placeholders = extract_placeholders(file)
                # Rewind file pointer after reading
                file.seek(0)
            except Exception as e:
                raise ValidationError({"file": f"Could not parse template file: {e}"})

        serializer.save(
            created_by=self.request.user,
            file_name=file.name if file else "",
            placeholders=placeholders,
        )

    def perform_destroy(self, instance):
        # Soft delete
        instance.is_active = False
        instance.save()

    @action(detail=True, methods=["post"])
    def duplicate(self, request, pk=None):
        """Create a copy of this template."""
        original = self.get_object()
        copy = DocumentTemplate.objects.create(
            name=f"{original.name} (Copy)",
            description=original.description,
            type=original.type,
            category=original.category,
            document_type=original.document_type,
            tags=original.tags,
            sections=original.sections,
            placeholders=original.placeholders,
            file=original.file,
            file_name=original.file_name,
            created_by=request.user,
        )
        serializer = self.get_serializer(copy)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=["get"])
    def placeholders(self, request, pk=None):
        template = self.get_object()
        return Response({"placeholders": template.placeholders})

    @action(detail=True, methods=["post"])
    def fill(self, request, pk=None):
        template = self.get_object()
        values = request.data.get("values", {})
        fmt = request.data.get("output_format", "pdf")
        title = request.data.get("title", template.name)
        type_id = request.data.get("document_type_id") or str(template.document_type_id or "")
        draft_from_template = bool(request.data.get("draft_from_template"))

        if not title or not title.strip():
            raise ValidationError({"title": "Document title is required."})

        if not template.document_type_id:
            raise ValidationError({"document_type_id": "This template is not attached to a document type."})
        if str(template.document_type_id) != str(type_id):
            raise ValidationError({"document_type_id": "Template does not belong to the selected document type."})

        if fmt not in ("pdf", "docx"):
            raise ValidationError({"output_format": "Must be 'pdf' or 'docx'."})

        # Validate required fields for built templates
        if template.type == "built" and not draft_from_template:
            all_fields = [
                f for section in template.sections
                for f in section.get("fields", [])
                if f.get("required") and f.get("type") not in ("divider", "heading")
            ]
            missing = [f["label"] for f in all_fields if not values.get(f["key"])]
            if missing:
                raise ValidationError(
                    {"values": f"Required fields missing: {', '.join(missing)}"}
                )

        try:
            doc = generate_document_from_template_sync(
                template, values, fmt, title.strip(), request.user, type_id
            )
        except Exception as e:
            raise ValidationError({"detail": f"Document generation failed: {e}"})

        # Increment use count atomically
        DocumentTemplate.objects.filter(id=template.id).update(
            use_count=F("use_count") + 1
        )
        TemplateUsage.objects.create(
            template=template,
            document=doc,
            used_by=request.user,
            values=values,
            output_format=fmt,
        )

        return Response({"document_id": str(doc.id)})

    @action(detail=True, methods=["get"])
    def usages(self, request, pk=None):
        """Return recent usage history for this template."""
        template = self.get_object()
        usages = template.usages.select_related("used_by", "document").order_by("-created_at")[:20]
        data = [
            {
                "id": str(u.id),
                "document_id": str(u.document_id),
                "document_title": u.document.title if u.document else "",
                "used_by": {
                    "id": str(u.used_by.id),
                    "full_name": u.used_by.get_full_name() or u.used_by.email,
                },
                "output_format": u.output_format,
                "created_at": u.created_at.isoformat(),
            }
            for u in usages
        ]
        return Response(data)
